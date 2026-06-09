"""
将 drama-copyright-check skill 输出的 Markdown 内容导出为 .docx 文件

用法：
    python export.py              ← 从剪贴板读取
    python export.py result.txt  ← 从文件读取

依赖：
    pip install python-docx pyperclip
"""

import re
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("缺少依赖，请先运行：pip install python-docx")
    sys.exit(1)


# ---------- 读取输入 ----------

def read_input() -> str:
    if len(sys.argv) >= 2:
        path = Path(sys.argv[1])
        if not path.exists():
            print(f"文件不存在：{path}")
            sys.exit(1)
        return path.read_text(encoding='utf-8')
    try:
        import pyperclip
        text = pyperclip.paste()
        if not text.strip():
            print("剪贴板为空，请先复制 skill 输出内容后再运行。")
            sys.exit(1)
        print("已从剪贴板读取内容。")
        return text
    except ImportError:
        print("缺少依赖，请运行：pip install pyperclip")
        sys.exit(1)


# ---------- 样式工具 ----------

def set_cell_background(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'AAAAAA')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_runs(paragraph, text: str, bold: bool = False, base_size: int = 10):
    """解析行内 **bold** 标记并添加 run。"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            run.bold = bold
        run.font.size = Pt(base_size)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


# ---------- Markdown 解析 ----------

def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """返回 (headers, data_rows)，自动跳过分隔行。"""
    headers = []
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        # 分隔行：所有单元格只含 -
        if all(re.fullmatch(r'-+', c) for c in cells if c):
            continue
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    return headers, rows


def add_table_to_doc(doc: Document, md_lines: list[str]):
    headers, rows = parse_table(md_lines)
    if not headers:
        return

    col_count = max(len(headers), max((len(r) for r in rows), default=0))
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'
    set_table_borders(table)

    # 设置列宽（均分页面宽度约 15cm）
    col_width = Cm(15.5 / col_count)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    # 表头行
    hdr_row = table.rows[0]
    for i, header in enumerate(headers[:col_count]):
        cell = hdr_row.cells[i]
        set_cell_background(cell, 'E8EFF7')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs(p, header, bold=True, base_size=9)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx in range(col_count):
            cell = row.cells[c_idx]
            text = row_data[c_idx] if c_idx < len(row_data) else ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs(p, text, base_size=9)
            # 隔行底色
            if r_idx % 2 == 1:
                set_cell_background(cell, 'F7F9FC')

    doc.add_paragraph()


# ---------- 文档构建 ----------

def build_doc(md_text: str) -> Document:
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 默认正文字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    lines = md_text.split('\n')
    i = 0
    table_buffer: list[str] = []

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            add_table_to_doc(doc, table_buffer)
            table_buffer = []

    while i < len(lines):
        line = lines[i].rstrip()

        # 表格行收集
        if line.startswith('|'):
            table_buffer.append(line)
            i += 1
            continue

        # 非表格行出现，先冲刷积累的表格
        flush_table()

        if line.startswith('# '):
            h = doc.add_heading(line[2:].strip(), level=1)
            h.runs[0].font.size = Pt(16)
        elif line.startswith('## '):
            h = doc.add_heading(line[3:].strip(), level=2)
            h.runs[0].font.size = Pt(13)
        elif line.startswith('### '):
            h = doc.add_heading(line[4:].strip(), level=3)
            h.runs[0].font.size = Pt(11)
        elif line.startswith('#### '):
            h = doc.add_heading(line[5:].strip(), level=4)
            h.runs[0].font.size = Pt(10)
        elif line == '---':
            doc.add_paragraph()
        elif line.strip():
            p = doc.add_paragraph()
            add_runs(p, line.strip(), base_size=10)
        # 空行跳过

        i += 1

    flush_table()
    return doc


# ---------- 主流程 ----------

def main():
    md_text = read_input()

    if len(sys.argv) >= 2:
        out_path = Path(sys.argv[1]).with_suffix('.docx')
    else:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        out_path = Path(f'版权分析报告_{timestamp}.docx')

    doc = build_doc(md_text)
    doc.save(str(out_path))
    print(f"已导出：{out_path.resolve()}")


if __name__ == '__main__':
    main()
